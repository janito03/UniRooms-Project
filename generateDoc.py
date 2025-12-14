# First, install the library: pip install python-docx

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for tag in ('top','left','bottom','right','insideH','insideV'):
        border = OxmlElement(f"w:{tag}")
        border.set(qn('w:val'), 'nil')
        borders.append(border)
    tblPr.append(borders)

def create_table(doc, headers, data):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Remove borders for single-column tables
    if len(headers) == 1:
        remove_table_borders(table)
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
    
    for row_data in data:
        row_cells = table.add_row().cells
        for i, item in enumerate(row_data):
            row_cells[i].text = str(item)
    doc.add_paragraph() # Spacer

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

doc.add_heading('Assignment 1, Project Summary', level=1)
create_table(doc, ['Course'], [['Web Technologies - 2025']])

doc.add_heading('Project Author', level=2)
create_table(doc, ['No', 'Name', 'FN'], [
    ['1', 'Student Name', '12345'],
    ['2', 'Partner Name', '67890']
])

doc.add_heading('Project Name', level=2)
create_table(doc, ['Project name'], [['UniRoom Scheduler']])

doc.add_heading('Short project description (Business needs and system features)', level=2)
desc_text = (
    "The UniRoom Scheduler allows students/teachers to view room availability and book resources. "
    "Developed using HTML/CSS/JS (Frontend) and Node.js/Express + MongoDB (Backend). "
    "Features include Base Schedule (immutable), Search/Filter, and User Roles: "
    "Anonymous, Student (Book max 2h), Teacher (Priority Override), Admin."
)
create_table(doc, ['Description'], [[desc_text]])

doc.add_heading('Main Use Cases / Scenarios', level=2)
create_table(doc, ['Use case name', 'Brief Descriptions', 'Actors Involved'], [
    ['Register/Login', 'Register as Student/Teacher. Login validates priority.', 'Anonymous'],
    ['Browse Schedule', 'View Base Schedule + Bookings.', 'All Users'],
    ['Filter & Search', 'Filter by Type (Lab/Lecture) or Time. Search by Room #.', 'Student, Teacher'],
    ['Reserve Room', 'Book free slot. Max duration: 2 hours.', 'Student, Teacher'],
    ['Teacher Override', 'Teacher cancels Student booking to claim slot.', 'Teacher'],
    ['Manage System', 'CRUD for Rooms, Users, and Base Schedule.', 'Administrator']
])

doc.add_heading('Main Views (SPA Frontend)', level=2)
create_table(doc, ['View name', 'Brief Descriptions', 'URI'], [
    ['Home/Login', 'Login/Register forms.', '/'],
    ['Dashboard', 'Interactive timetable with filters.', '/dashboard'],
    ['Room Details', 'Room features and weekly view.', '/rooms/{id}'],
    ['My Bookings', 'User reservations list.', '/my-bookings'],
    ['Admin Panel', 'Manage Users/Rooms/Base Schedule.', '/admin']
])

doc.add_heading('API Resources (Node.js Backend)', level=2)
create_table(doc, ['Resource', 'Brief Descriptions', 'URI'], [
    ['Auth', 'Login/Register.', '/api/auth'],
    ['Users', 'Manage users/roles.', '/api/users'],
    ['Rooms', 'CRUD Rooms.', '/api/rooms'],
    ['Schedule', 'GET combined schedule.', '/api/schedule'],
    ['Bookings', 'POST reservation (checks 2h limit).', '/api/bookings'],
    ['Override', 'POST teacher override logic.', '/api/bookings/override']
])

doc.save('UniRoom_Project_Specs.docx')
print("Document created successfully.")