from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_table(doc, headers, data):
    # Create table with headers
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header styling
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    # Data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for i, item in enumerate(row_data):
            row_cells[i].text = str(item)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    doc.add_paragraph()  # Spacer

# Initialize Document
doc = Document()

# --- Title ---
title = doc.add_heading('University Room Booking - Database Configuration', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# --- 1. Database Overview ---
doc.add_heading('1. Database Overview', level=1)
headers_overview = ['Configuration', 'Details']
data_overview = [
    ['Database System', 'MongoDB'],
    ['ODM/Framework', 'Mongoose (Node.js)'],
    ['Architecture', 'Document-Oriented (Collections)']
]
create_table(doc, headers_overview, data_overview)

# --- 2. Collection Schemas ---
doc.add_heading('2. Collection Schemas', level=1)

# A. Users Collection
doc.add_heading('A. Users Collection', level=2)
headers_users = ['Field', 'Type', 'Constraints', 'Description']
data_users = [
    ['_id', 'ObjectId', 'Auto-generated', 'Unique ID'],
    ['username', 'String', 'Unique, Required', 'Login handle'],
    ['password_hash', 'String', 'Required', 'Encrypted password'],
    ['email', 'String', 'Required', 'Contact email'],
    ['role', 'String', "Enum: ['student',...]", 'Access rights'],
    ['createdAt', 'Date', 'Default: Date.now', 'Creation time']
]
create_table(doc, headers_users, data_users)

# B. Rooms Collection
doc.add_heading('B. Rooms Collection', level=2)
headers_rooms = ['Field', 'Type', 'Constraints', 'Description']
data_rooms = [
    ['id', 'ObjectId', 'Auto-generated', 'Unique ID'],
    ['roomNumber', 'String', 'Unique, Required', 'e.g., "101B"'],
    ['capacity', 'Integer', 'Required', 'Max people'],
    ['type', 'String', "Enum: ['normal'...]", 'Filter category'],
    ['features', 'Array', 'Optional', 'e.g., Projector']
]
create_table(doc, headers_rooms, data_rooms)

# C. BaseSchedule Collection
doc.add_heading('C. BaseSchedule Collection (Official Timetable)', level=2)
headers_schedule = ['Field', 'Type', 'Constraints', 'Description']
data_schedule = [
    ['id', 'ObjectId', 'Auto-generated', 'Unique ID'],
    ['room_id', 'ObjectId', 'Ref: Rooms', 'Room Link'],
    ['subject', 'String', 'Required', 'Class Name'],
    ['dayOfWeek', 'String', "Enum: ['Mon'...]", 'Repeats on'],
    ['startTime', 'String', '24h format', 'Start (08:30)'],
    ['endTime', 'String', '24h format', 'End (10:00)'],
    ['semester', 'String', 'Required', 'Current Term']
]
create_table(doc, headers_schedule, data_schedule)

# D. Bookings Collection
doc.add_heading('D. Bookings Collection (Dynamic)', level=2)
headers_bookings = ['Field', 'Type', 'Constraints', 'Description']
data_bookings = [
    ['_id', 'Objectid', 'Auto-generated', 'Unique ID'],
    ['room_id', 'ObjectId', 'Ref: Rooms', 'Room Link'],
    ['user_id', 'Objectid', 'Ref: Users', 'User Link'],
    ['start Time', 'Date', 'Required', 'ISO Timestamp'],
    ['endTime', 'Date', 'Required', 'ISO Timestamp'],
    ['status', 'String', "Enum: ['confirmed'...]", 'Booking State'],
    ['overriddenBy', 'ObjectId', 'Ref: Users', 'Teacher Override']
]
create_table(doc, headers_bookings, data_bookings)

# --- 3. Priority Logic ---
doc.add_page_break()
doc.add_heading('3. Priority Logic (Teacher Override)', level=1)

doc.add_paragraph('Conflict Detection:', style='List Bullet')
p = doc.add_paragraph('(newStart < existing End) AND (newEnd > existingStart)')
p.paragraph_format.left_indent = Inches(0.5)

doc.add_paragraph('Booking Request Flow:', style='List Bullet')
flow_steps = [
    '1. Check Base Schedule for official classes -> Block if conflict',
    '2. Check Existing Bookings:',
    '    a. If empty -> Allow',
    '    b. If booked by teacher -> Reject (Teachers have priority)',
    '    c. If booked by student & requester is teacher -> Override/Cancel',
    '    d. If booked by student & requester is student -> Reject (max 2 hours)'
]
for step in flow_steps:
    p = doc.add_paragraph(step)
    p.paragraph_format.left_indent = Inches(0.5)

# --- 4. Database Indexing ---
doc.add_heading('4. Database Indexing & Optimization', level=1)

doc.add_heading('Room Indexes:', level=2)
doc.add_paragraph('- { roomNumber: 1 } - Fast lookup by room code', style='List Bullet')
doc.add_paragraph('- { type: 1, capacity: 1 } - Filter by attributes', style='List Bullet')

doc.add_heading('Booking Indexes:', level=2)
doc.add_paragraph('- { room_id: 1, startTime: 1, endTime: 1 } - Conflict detection', style='List Bullet')
doc.add_paragraph('- { user_id: 1, startTime: -1 } - User history & 2-hour check', style='List Bullet')

# Save
doc.save('University_DB_Config.docx')
print("Document generated successfully: University_DB_Config.docx")